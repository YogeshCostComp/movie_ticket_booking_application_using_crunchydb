"""
Generate SRE Agent Architecture & Flow Documentation (Word .docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = "SRE_Agent_Architecture_Flow.docx"

def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_style(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return h

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "1F2937")

    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def create_doc():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SRE Agent Orchestrator")
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Architecture & Flow Documentation")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Ephemeral Agent Architecture with MCP Server Integration")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Technologies: FastAPI • Anthropic Claude • MCP Protocol • WebSocket\n"
                        "Deployment: Render / Docker / IBM Cloud Code Engine\n"
                        "Author: SRE Team | February 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Architecture Overview",
        "3. Component Details",
        "   3.1 SRE Agent Orchestrator (FastAPI)",
        "   3.2 LLM Brain (Anthropic Claude)",
        "   3.3 MCP Client",
        "   3.4 Ephemeral Agents",
        "   3.5 Agent Registry",
        "   3.6 Frontend (Chat + Pipeline UI)",
        "4. Ephemeral Agent Lifecycle",
        "5. Request Flow — Step by Step",
        "6. Architecture Diagrams (ASCII)",
        "7. Agent Types Reference",
        "8. API Endpoints",
        "9. Deployment Guide (Render)",
        "10. Environment Variables",
        "11. File Structure",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "1. Executive Summary", level=1)
    doc.add_paragraph(
        "The SRE Agent Orchestrator is an intelligent, AI-powered Site Reliability Engineering (SRE) "
        "assistant that uses an ephemeral agent architecture to handle operational tasks. When a user "
        "asks a question (e.g., \"check error logs\", \"is the app healthy?\"), the system:"
    )
    bullets = [
        "Analyzes the intent using Anthropic Claude LLM",
        "Spins up a specialized ephemeral agent (e.g., Log Agent, Health Agent)",
        "The agent connects to the MCP Server on IBM Cloud Code Engine",
        "Retrieves real-time data from the production application",
        "Formats the response using Claude AI and streams it back to the user",
        "The agent is destroyed after a configurable cooldown period (default: 120 seconds)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph(
        "\nThe key innovation is that agents are truly ephemeral — they are Python objects that are "
        "created on-demand, execute their task, and are garbage-collected. The system provides full "
        "lifecycle proof via the Agent Registry, which tracks object IDs, memory addresses, PIDs, "
        "thread info, and a complete audit trail of every event."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 2. ARCHITECTURE OVERVIEW
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "2. Architecture Overview", level=1)
    doc.add_paragraph(
        "The system follows a layered architecture with clear separation of concerns:"
    )

    # ASCII Architecture Diagram
    arch_diagram = """
┌─────────────────────────────────────────────────────────────────────────┐
│                        USER / BROWSER                                   │
│                    Chat UI + Pipeline Panel                              │
│                  (HTML/CSS/JS + WebSocket)                               │
└─────────────────────┬───────────────────┬───────────────────────────────┘
                      │ WebSocket         │ REST API
                      ▼                   ▼
┌─────────────────────────────────────────────────────────────────────────┐
│                   SRE AGENT ORCHESTRATOR                                │
│                     (FastAPI Server)                                     │
│                                                                         │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────────────────────┐  │
│  │  LLM Brain   │  │ Agent Router │  │     Agent Registry           │  │
│  │  (Claude AI) │  │  (Spawner)   │  │  (Lifecycle Tracker)         │  │
│  └──────┬───────┘  └──────┬───────┘  └──────────────────────────────┘  │
│         │                 │                                              │
│         ▼                 ▼                                              │
│  ┌─────────────────────────────────────────────────────────────────┐    │
│  │            EPHEMERAL AGENT POOL (Created on-demand)             │    │
│  │                                                                 │    │
│  │  📋 Log Agent    🏥 Health Agent    📡 Monitoring Agent         │    │
│  │  📕 Runbook Agent   🔗 Trace Agent    📊 Dashboard Agent       │    │
│  │  🚀 Deployment Agent                                           │    │
│  │                                                                 │    │
│  │  Each agent: Created → Executes → Cooldown → Destroyed          │    │
│  └─────────────────────────┬───────────────────────────────────────┘    │
│                            │                                            │
└────────────────────────────┼────────────────────────────────────────────┘
                             │ HTTP / JSON-RPC 2.0
                             ▼
┌─────────────────────────────────────────────────────────────────────────┐
│                    MCP SERVER (IBM Cloud Code Engine)                    │
│                   Flask App — 30+ SRE Tools                             │
│                                                                         │
│  Health Tools │ Log Tools │ Trace Tools │ Monitoring │ Deployment       │
│  check_app    │ error_logs│ get_traces  │ start/stop │ restart_app      │
│  check_db     │ app_logs  │ trace_detail│ status     │ app_status       │
│  system_status│ query_logs│ summary     │ runbook    │ deploy_history   │
└─────────────────────────────┬───────────────────────────────────────────┘
                              │
                              ▼
┌─────────────────────────────────────────────────────────────────────────┐
│              MOVIE TICKET BOOKING APPLICATION                           │
│              Flask + PostgreSQL (Neon Serverless)                        │
│              IBM Cloud Code Engine                                       │
└─────────────────────────────────────────────────────────────────────────┘
"""
    p = doc.add_paragraph()
    run = p.add_run(arch_diagram)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 3. COMPONENT DETAILS
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "3. Component Details", level=1)

    # 3.1
    add_heading_with_style(doc, "3.1 SRE Agent Orchestrator (FastAPI)", level=2)
    doc.add_paragraph("File: sre-agent/app.py")
    doc.add_paragraph(
        "The orchestrator is the central hub of the system. It is a FastAPI application that:"
    )
    items = [
        "Serves the web UI at / (Chat + Pipeline panel)",
        "Accepts WebSocket connections at /ws for real-time communication",
        "Routes user queries to the LLM Brain for intent classification",
        "Spawns the appropriate ephemeral agent based on classified intent",
        "Broadcasts pipeline events to all connected clients in real-time",
        "Manages agent cooldown and delayed destruction (configurable, default 120s)",
        "Provides REST API endpoints for agent inspection and history",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # 3.2
    add_heading_with_style(doc, "3.2 LLM Brain (Anthropic Claude)", level=2)
    doc.add_paragraph("File: sre-agent/llm_brain.py")
    doc.add_paragraph(
        "The LLM Brain wraps the Anthropic Claude API (claude-sonnet-4-20250514) and provides two "
        "core functions:"
    )
    doc.add_paragraph(
        "1. Intent Classification (classify_intent): Takes a natural language user query and "
        "returns a JSON object with: agent name, action, parameters, and reasoning. This decides "
        "which ephemeral agent to spawn."
    )
    doc.add_paragraph(
        "2. Response Formatting (format_response): Takes raw MCP tool output (JSON data) and "
        "transforms it into a human-readable markdown response with status indicators, tables, "
        "and actionable insights."
    )

    # 3.3
    add_heading_with_style(doc, "3.3 MCP Client", level=2)
    doc.add_paragraph("File: sre-agent/mcp_client.py")
    doc.add_paragraph(
        "The MCP Client communicates with the SRE MCP Server hosted on IBM Cloud Code Engine. "
        "It supports two communication protocols:"
    )
    doc.add_paragraph("• JSON-RPC 2.0 (POST /mcp) — For standard MCP tool invocations via tools/call and tools/list")
    doc.add_paragraph("• REST (POST /tools/<tool_name>) — For direct tool endpoint calls")
    doc.add_paragraph(
        "\nAuthentication is via X-API-Key header. The client has 30+ convenience methods "
        "mapping to specific MCP tools (get_error_logs, check_app_health, get_recent_traces, etc.)."
    )

    # 3.4
    add_heading_with_style(doc, "3.4 Ephemeral Agents", level=2)
    doc.add_paragraph("Files: sre-agent/agents/*.py")
    doc.add_paragraph(
        "Ephemeral agents are the core innovation. Each agent is a Python object that:"
    )
    items = [
        "Inherits from BaseAgent (abstract class in base_agent.py)",
        "Gets a unique agent_id (UUID-based) on creation",
        "Registers itself in the Agent Registry on __init__",
        "Has an event_callback to emit pipeline events (visible in the UI)",
        "Executes its specific task by calling MCP tools via the MCP Client",
        "Is held in a cooldown pool for N seconds (for demo/inspection)",
        "Is then deregistered and garbage-collected (truly destroyed)",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "\nWhy they are \"real agents\" and not just functions:"
    )
    items_why = [
        "Autonomy: They decide which MCP tools to call based on the action",
        "Identity: Each has a unique ID, its own memory address, its own lifecycle",
        "Observability: They emit events, record audit trails, are inspectable while alive",
        "Lifecycle: They are born (created), work (execute), rest (cooldown), and die (destroyed)",
        "Ephemerality: Unlike long-running services, they exist ONLY for the duration of the task",
    ]
    for item in items_why:
        doc.add_paragraph(item, style='List Bullet')

    # 3.5
    add_heading_with_style(doc, "3.5 Agent Registry", level=2)
    doc.add_paragraph("File: sre-agent/agent_registry.py")
    doc.add_paragraph(
        "The Agent Registry is a singleton that provides proof of agent lifecycle. "
        "It tracks every agent with:"
    )
    add_styled_table(doc,
        ["Field", "Description", "Example"],
        [
            ["agent_id", "Unique UUID-based identifier", "agent-fdf87233"],
            ["python_object_id", "Python id() — memory address", "0x181CD5796A0"],
            ["python_class", "Python class name", "LogAgent"],
            ["process_id", "OS Process ID", "5304"],
            ["thread_id / thread_name", "Thread information", "MainThread (12345)"],
            ["created_at", "ISO timestamp of creation", "2026-02-19T12:19:05"],
            ["completed_at", "ISO timestamp of completion", "2026-02-19T12:19:17"],
            ["duration_seconds", "Total execution time", "12.3s"],
            ["status", "Current lifecycle state", "active / executing / destroyed"],
            ["events[]", "Full audit trail of pipeline events", "Array of event objects"],
        ]
    )

    # 3.6
    doc.add_paragraph()
    add_heading_with_style(doc, "3.6 Frontend (Chat + Pipeline UI)", level=2)
    doc.add_paragraph("Files: sre-agent/templates/index.html, sre-agent/static/js/app.js, sre-agent/static/css/style.css")
    doc.add_paragraph(
        "The frontend is a single-page application with two panels:"
    )
    doc.add_paragraph("Left Panel — Chat Window: Users type natural language queries. Agent responses are rendered as markdown.")
    doc.add_paragraph("Right Panel — Agent Pipeline: Shows real-time pipeline steps as agents execute, with the active agent card, step indicators, and a history section.")
    doc.add_paragraph(
        "\nKey features: WebSocket for real-time updates, Inspect Drawer (slide-in overlay showing full agent lifecycle proof), "
        "session persistence via sessionStorage (chat and pipeline survive page refreshes), hint chips for quick queries."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 4. EPHEMERAL AGENT LIFECYCLE
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "4. Ephemeral Agent Lifecycle", level=1)

    lifecycle_diagram = """
    ┌───────────────────────────────────────────────────────────────┐
    │                  AGENT LIFECYCLE STATES                       │
    └───────────────────────────────────────────────────────────────┘

    ┌──────────┐    ┌───────────┐    ┌───────────┐    ┌───────────┐
    │          │    │           │    │           │    │           │
    │ CREATED  │───▶│ EXECUTING │───▶│ COOLDOWN  │───▶│ DESTROYED │
    │          │    │           │    │           │    │           │
    └──────────┘    └───────────┘    └───────────┘    └───────────┘
         │               │               │                 │
    Registered in   Calls MCP tools   Agent alive but   Deregistered,
    Agent Registry  Emits pipeline    idle. Inspectable  Python GC'd
    Gets unique ID  events to UI      via /inspect/{id}  Memory freed
    Memory allocated                  (120s default)

    Timeline Example:
    ─────────────────────────────────────────────────────────────────
    t=0s     Agent Created (agent-fdf87233, obj@0x181CD5796A0)
    t=0.1s   Registered in AgentRegistry
    t=0.2s   Intent classified → action: get_error_logs
    t=0.5s   Connecting to MCP Server
    t=2.0s   MCP tool call: get_error_logs(hours=24)
    t=8.5s   Data received, processing results
    t=10.0s  Claude AI formatting response
    t=12.3s  Response sent to user — Agent COMPLETED
    t=12.3s  ──── COOLDOWN STARTS (120 seconds) ────
    t=132.3s Agent DESTROYED — deregistered, garbage-collected
"""
    p = doc.add_paragraph()
    run = p.add_run(lifecycle_diagram)
    run.font.name = 'Consolas'
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 5. REQUEST FLOW
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "5. Request Flow — Step by Step", level=1)
    doc.add_paragraph(
        "This section walks through the complete flow when a user asks: \"Check error logs in the last 24 hours\""
    )

    flow_steps = [
        ("Step 1: User Input",
         "User types \"Check error logs in the last 24 hours\" in the chat window.\n"
         "The message is sent over WebSocket to the FastAPI orchestrator at /ws."),
        ("Step 2: Intent Classification (LLM Brain)",
         "The orchestrator sends the user message to Claude AI (classify_intent).\n"
         "Claude returns JSON: { \"agent\": \"log_agent\", \"action\": \"get_error_logs\", "
         "\"params\": { \"hours\": 24, \"limit\": 100 }, \"reasoning\": \"User wants error logs\" }"),
        ("Step 3: Pipeline Event — Intent Classified",
         "A pipeline event is broadcast to all connected WebSocket clients.\n"
         "The UI shows: ✅ Intent classified — Agent: log_agent | Action: get_error_logs"),
        ("Step 4: Ephemeral Agent Creation",
         "The orchestrator looks up log_agent in AGENT_CLASSES dict.\n"
         "A new LogAgent object is instantiated: agent = LogAgent(mcp_client=mcp, event_callback=callback)\n"
         "During __init__, the agent: gets a UUID (agent-fdf87233), registers in AgentRegistry, "
         "records its Python object ID, PID, thread info."),
        ("Step 5: Agent Execution",
         "agent.run(\"get_error_logs\", {\"hours\": 24, \"limit\": 100}) is called.\n"
         "The agent calls self.mcp.get_error_logs(24, 100) which sends HTTP POST to the MCP Server.\n"
         "The MCP Server queries the actual application logs and returns the data.\n"
         "The agent emits pipeline events at each step (visible in real-time in the UI)."),
        ("Step 6: Pipeline Event — Agent Available for Inspection",
         "The orchestrator emits an event with inspect_url: /inspect/{agent_id}.\n"
         "The UI shows an \"Inspect Live Agent\" button and a cooldown timer.\n"
         "The agent object is placed in the _cooldown_agents dict to keep a Python reference alive."),
        ("Step 7: Response Formatting (LLM Brain)",
         "The raw MCP data (JSON) is sent to Claude AI (format_response).\n"
         "Claude transforms it into a clean markdown report with ✅/❌ status indicators, tables, "
         "and actionable recommendations."),
        ("Step 8: Chat Response",
         "The formatted markdown is sent to the user's chat via WebSocket.\n"
         "The UI renders it using marked.js (markdown → HTML)."),
        ("Step 9: Cooldown Period (120 seconds)",
         "The agent remains alive and inspectable. The user can click \"Inspect Live Agent\" to see:\n"
         "• Agent ID, Python class, memory address (object ID), PID, thread info\n"
         "• Creation/completion timestamps, duration, action executed\n"
         "• Full lifecycle audit trail of every pipeline event\n"
         "A countdown timer shows seconds until auto-destruction."),
        ("Step 10: Agent Destruction",
         "After the cooldown expires, delayed_agent_destruction() runs:\n"
         "• registry.deregister(agent, result) — moves agent to completed list\n"
         "• _cooldown_agents.pop(agent_id) — removes Python reference\n"
         "• Python garbage collector frees the memory\n"
         "• A final pipeline event is broadcast: 🗑️ Destroying Log Analysis Agent\n"
         "• The agent card in the UI shows status: DESTROYED"),
    ]

    for title, desc in flow_steps:
        add_heading_with_style(doc, title, level=3)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 6. ARCHITECTURE DIAGRAMS
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "6. Architecture Diagrams", level=1)

    add_heading_with_style(doc, "6.1 Sequence Diagram — User Query Flow", level=2)
    seq_diagram = """
    User        Orchestrator     LLM Brain       Agent         MCP Server      App
     │              │                │              │               │            │
     │──message────▶│                │              │               │            │
     │              │──classify──────▶│              │               │            │
     │              │◀──intent json──│              │               │            │
     │              │                │              │               │            │
     │              │──create agent──────────────▶  │               │            │
     │              │  (LogAgent.__init__)          │               │            │
     │              │                │              │               │            │
     │◀──pipeline──│                │              │               │            │
     │   events     │──run(action,params)────────▶  │               │            │
     │              │                │              │──HTTP POST───▶│            │
     │              │                │              │               │──query────▶│
     │              │                │              │               │◀──data─────│
     │              │                │              │◀──JSON────────│            │
     │◀──pipeline──│◀──result───────────────────── │               │            │
     │   events     │                │              │               │            │
     │              │──format────────▶│              │               │            │
     │              │◀──markdown─────│              │               │            │
     │              │                │              │               │            │
     │◀──chat msg──│                │              │               │            │
     │              │                │  ┌──120s──┐  │               │            │
     │◀──inspect───│                │  │COOLDOWN│  │               │            │
     │   link       │                │  └───┬────┘  │               │            │
     │              │──deregister────────────▶ ✝    │               │            │
     │◀──destroyed─│  (agent GC'd)         DEAD    │               │            │
     │              │                │              │               │            │
"""
    p = doc.add_paragraph()
    run = p.add_run(seq_diagram)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    add_heading_with_style(doc, "6.2 Agent Spawn & Destroy Cycle", level=2)
    spawn_diagram = """
    ┌─────────── ORCHESTRATOR ──────────────────────────────────────────┐
    │                                                                    │
    │  User Query ──▶ classify_intent() ──▶ "log_agent"                 │
    │                                                                    │
    │  AGENT_CLASSES["log_agent"] = LogAgent                             │
    │                                                                    │
    │  agent = LogAgent(mcp_client=mcp, event_callback=cb)               │
    │           │                                                        │
    │           ├─▶ self.agent_id = "agent-fdf87233"                     │
    │           ├─▶ registry.register(self)  ←── PROOF OF CREATION       │
    │           │      ├─ python_object_id = id(agent) = 0x181CD5796A0   │
    │           │      ├─ process_id = 5304                              │
    │           │      ├─ thread = MainThread                            │
    │           │      └─ created_at = 2026-02-19T12:19:05               │
    │           │                                                        │
    │  result = await agent.run("get_error_logs", params)                │
    │           │                                                        │
    │           ├─▶ self.mcp.get_error_logs(24, 100)  ──▶ MCP Server     │
    │           ├─▶ emit("📋 Log data retrieved")                        │
    │           └─▶ return {status: "success", data: {...}}              │
    │                                                                    │
    │  _cooldown_agents[agent_id] = agent  ←── KEEP ALIVE               │
    │  asyncio.create_task(delayed_agent_destruction(agent, ...))        │
    │           │                                                        │
    │           ├─▶ await asyncio.sleep(120)  ←── COOLDOWN               │
    │           ├─▶ registry.deregister(agent, result)                   │
    │           │      ├─ completed_at = 2026-02-19T12:21:05             │
    │           │      ├─ duration_seconds = 120.0                       │
    │           │      └─ status = "destroyed"                           │
    │           └─▶ _cooldown_agents.pop(agent_id)  ←── REF REMOVED      │
    │                       │                                            │
    │                       ▼                                            │
    │              Python GC collects agent object                       │
    │              Memory at 0x181CD5796A0 is freed                      │
    │                                                                    │
    └────────────────────────────────────────────────────────────────────┘
"""
    p = doc.add_paragraph()
    run = p.add_run(spawn_diagram)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 7. AGENT TYPES REFERENCE
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "7. Agent Types Reference", level=1)

    add_styled_table(doc,
        ["Agent", "Icon", "File", "Actions", "Description"],
        [
            ["Log Agent", "📋", "log_agent.py", "get_error_logs, get_recent_logs, get_app_logs, get_platform_logs, query_logs", "Analyzes application and platform logs, searches for error patterns"],
            ["Health Agent", "🏥", "health_agent.py", "check_app_health, check_database_health, get_system_status, check_all", "Runs health checks on app, database, and full system"],
            ["Monitoring Agent", "📡", "monitoring_agent.py", "start, stop, status", "Starts/stops continuous monitoring with configurable interval"],
            ["Runbook Agent", "📕", "runbook_agent.py", "start, stop, status", "Automated runbook with auto-restart on detected errors"],
            ["Trace Agent", "🔗", "trace_agent.py", "get_recent_traces, get_trace_details, get_trace_summary", "Analyzes distributed request traces, finds slow endpoints"],
            ["Dashboard Agent", "📊", "dashboard_agent.py", "get_dashboard, get_response_times, get_failure_analysis", "Builds SRE dashboard with golden signals (latency, traffic, errors, saturation)"],
            ["Deployment Agent", "🚀", "deployment_agent.py", "get_deployment_history, get_app_status, restart_app, stop_app, start_app", "Manages application lifecycle — deploy, restart, stop, start"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 8. API ENDPOINTS
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "8. API Endpoints", level=1)

    add_styled_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/", "Serves the main Chat + Pipeline UI"],
            ["GET", "/health", "Health check — reports MCP server reachability"],
            ["WebSocket", "/ws", "Real-time bidirectional communication for chat and pipeline events"],
            ["POST", "/api/query", "REST alternative to WebSocket for single queries"],
            ["GET", "/api/agents", "List all available agent types"],
            ["GET", "/api/history", "Get recent agent run history"],
            ["GET", "/api/agents/active", "List currently alive agents (in cooldown or executing)"],
            ["GET", "/api/agents/completed", "List destroyed agents with full audit trail"],
            ["GET", "/api/agents/stats", "High-level stats: total created, destroyed, active count"],
            ["GET", "/api/agents/{agent_id}", "Inspect a specific agent — full lifecycle with all events"],
            ["GET", "/inspect/{agent_id}", "Agent inspection dashboard page (HTML)"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 9. DEPLOYMENT GUIDE (RENDER)
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "9. Deployment Guide (Render)", level=1)

    add_heading_with_style(doc, "Option A: Blueprint Deployment (Recommended)", level=2)
    steps_a = [
        "1. Push code to GitHub (already done — repo: YogeshCostComp/movie_ticket_booking_application_using_crunchydb)",
        "2. Go to https://dashboard.render.com → New → Blueprint",
        "3. Connect your GitHub repository",
        "4. Render will detect the render.yaml file in sre-agent/ folder",
        "5. It will auto-configure: Python runtime, build command, start command, root directory",
        "6. Set the ANTHROPIC_API_KEY environment variable (marked as sync: false in render.yaml for security)",
        "7. Click 'Apply' — Render will build and deploy automatically",
    ]
    for s in steps_a:
        doc.add_paragraph(s)

    add_heading_with_style(doc, "Option B: Manual Web Service", level=2)
    steps_b = [
        "1. Go to https://dashboard.render.com → New → Web Service",
        "2. Connect GitHub repo: YogeshCostComp/movie_ticket_booking_application_using_crunchydb",
        "3. Set Root Directory: sre-agent",
        "4. Set Build Command: pip install -r requirements.txt",
        "5. Set Start Command: uvicorn app:app --host 0.0.0.0 --port $PORT",
        "6. Set Environment: Python 3",
        "7. Add Environment Variables (see Section 10 below)",
        "8. Click 'Create Web Service'",
    ]
    for s in steps_b:
        doc.add_paragraph(s)

    # ═══════════════════════════════════════════════════════════════
    # 10. ENVIRONMENT VARIABLES
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "10. Environment Variables", level=1)

    add_styled_table(doc,
        ["Variable", "Required", "Default", "Description"],
        [
            ["ANTHROPIC_API_KEY", "✅ Yes", "(none)", "Anthropic API key for Claude AI"],
            ["MCP_SERVER_URL", "✅ Yes", "https://sre-mcp-server.260m2gai7zqb...", "URL of the SRE MCP Server on IBM Cloud"],
            ["MCP_API_KEY", "No", "sre-mcp-secret-key-2026", "API key for MCP Server authentication"],
            ["AGENT_COOLDOWN_SECONDS", "No", "120", "Seconds to keep agent alive before destruction"],
            ["PORT", "No", "8000", "Server port (Render sets this automatically)"],
            ["PYTHON_VERSION", "No", "3.11.0", "Python version for Render"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 11. FILE STRUCTURE
    # ═══════════════════════════════════════════════════════════════
    add_heading_with_style(doc, "11. File Structure", level=1)

    file_tree = """
sre-agent/
├── app.py                    # FastAPI orchestrator — main entry point
├── llm_brain.py              # Anthropic Claude integration (intent + formatting)
├── mcp_client.py             # HTTP client for MCP Server (JSON-RPC + REST)
├── agent_registry.py         # Singleton lifecycle tracker for all agents
├── agents/
│   ├── __init__.py
│   ├── base_agent.py         # Abstract base class for all ephemeral agents
│   ├── log_agent.py          # 📋 Log analysis agent
│   ├── health_agent.py       # 🏥 Health check agent
│   ├── monitoring_agent.py   # 📡 Monitoring control agent
│   ├── runbook_agent.py      # 📕 Runbook automation agent
│   ├── trace_agent.py        # 🔗 Trace analysis agent
│   ├── dashboard_agent.py    # 📊 SRE dashboard agent
│   └── deployment_agent.py   # 🚀 Deployment management agent
├── templates/
│   ├── index.html            # Main UI — chat + pipeline + inspect overlay
│   └── inspect.html          # Standalone agent inspection dashboard
├── static/
│   ├── css/
│   │   └── style.css         # Full dark-theme styles (~1000 lines)
│   └── js/
│       └── app.js            # WebSocket client, pipeline, inspect panel (~560 lines)
├── Dockerfile                # Docker image definition (Python 3.11 + uvicorn)
├── render.yaml               # Render blueprint for automated deployment
└── requirements.txt          # Python dependencies
"""
    p = doc.add_paragraph()
    run = p.add_run(file_tree)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # ═══════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════
    filepath = os.path.join(
        r"C:\Users\yokrishn\.vscode\movie_ticket_booking_application_using_crunchydb-master",
        OUTPUT
    )
    doc.save(filepath)
    print(f"✅ Document saved: {filepath}")


if __name__ == "__main__":
    create_doc()
